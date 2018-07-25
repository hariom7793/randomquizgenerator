#! python3
# randomQuizGenerator.py - Creates quizzes with questions and answers in
# random order, along with the answer key.

import random, openpyxl, docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class RandomQuiz:

    @staticmethod
    def open_files():
        return docx.Document(), docx.Document()

    def write_file_header(self, file_num, q_paper, ans_paper):
        q_paper.add_paragraph('Name:\n\nDate:\n\nPeriod:\n')
        q_paper.add_paragraph(self.quiz_title + ' (Paper %s)' %(file_num + 1)).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        ans_paper.add_paragraph(self.quiz_title + ' (Paper %s) Answer Key' %(file_num + 1)).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @staticmethod
    def break_page(q_paper, ans_paper):
        temp_q_para = q_paper.add_paragraph()
        temp_q_para.add_run()
        temp_q_para.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
        temp_a_para = ans_paper.add_paragraph()
        temp_a_para.add_run()
        temp_a_para.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

    @staticmethod
    def close_files(q_paper, ans_paper):
        q_paper.save('QuestionPapers.docx')
        ans_paper.save('QuestionPaperAnswers.docx')

    def gen_quest_ans_dict(self, filename):
        try:
            wb = openpyxl.load_workbook(filename)
            sheet = wb.active
            for row in range(1, sheet.max_row + 1):
                question = sheet['A' + str(row)].value
                answer = sheet['B' + str(row)].value
                if question is None or answer is None:
                    self.err_title = "Invalid Data - Excel"
                    self.err_msg = "Empty Question/Answer in given excel sheet."
                    break
                self.quest_ans_dict[question] = answer
                self.num_of_questions = sheet.max_row
        except:
            self.err_title = "Invalid Excel Sheet"
            self.err_msg = "Please select a valid excel sheet."

    def check_valid_inputs(self):
        if self.err_msg == "":
            if self.quiz_title == 0:
                self.err_title = "Invalid Title"
                self.err_msg = "Paper Title can not be empty."
            elif self.num_of_output_files == 0:
                self.err_title = "Invalid Question Paper count"
                self.err_msg = "Number of question papers can not be 0."
            elif self.num_of_questions == 0:
                self.err_title = "Invalid Question/Answer Excel"
                self.err_msg = "Empty Excel Sheet. Please check!"

    # Call below method to generate files
    def generate_random_quiz(self):

        if self.err_msg != "":
            return False, self.err_title, self.err_msg

        # Create the quiz and answer key files.
        quiz_file, answer_key_file = self.open_files()

        # Write question and answers.
        for quizNum in range(self.num_of_output_files):

            # Write out the header for the quiz.
            self.write_file_header(quizNum, quiz_file, answer_key_file)

            # Shuffle the order of the states.
            questions = list(self.quest_ans_dict.keys())  # get all questions in a list
            random.shuffle(questions)  # randomize the order of the questions

            answer_file_para = answer_key_file.add_paragraph()
            # Loop through all questions, making a question for each.
            for questionNum in range(self.num_of_questions):

                # Get right and wrong answers.
                correct_answer = self.quest_ans_dict[questions[questionNum]]
                wrong_answers = list(self.quest_ans_dict.values())  # get a complete list of answers
                del wrong_answers[wrong_answers.index(correct_answer)]  # remove the right answer
                wrong_answers = random.sample(wrong_answers, 3)  # pick 3 random ones

                answer_options = wrong_answers + [correct_answer]
                random.shuffle(answer_options)  # randomize the order of the answers

                # Write the question and answer options to the quiz file.
                quiz_file_para = quiz_file.add_paragraph('%s. %s\n' % (questionNum + 1, questions[questionNum]))
                for i in range(4):
                    quiz_file_para.add_run('  %s. %s\n' % ('ABCD'[i], answer_options[i]))

                # Write out the answer key to a file.
                answer_file_para.add_run('%s. %s\n' % (questionNum + 1, 'ABCD'[answer_options.index(correct_answer)]))

            # Switch to next page in quiz and answer key files to print next paper.
            self.break_page(quiz_file, answer_key_file)

        # Close the quiz and answer key files.
        self.close_files(quiz_file, answer_key_file)

        return True, "Successfully Generated Question Papers", "Question papers and answer key generated successfully."

    def __init__(self, quiz_title, quest_ans_file, num_of_output_files):

        self.quiz_title = quiz_title
        self.num_of_questions = 0
        self.quest_ans_dict = {}
        self.num_of_output_files = num_of_output_files

        self.err_title = ""
        self.err_msg = ""

        self.gen_quest_ans_dict(quest_ans_file)

        self.check_valid_inputs()

